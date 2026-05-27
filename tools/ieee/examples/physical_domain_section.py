# -*- coding: utf-8 -*-
"""
Example: Generate a standalone IEEE-formatted section (III.E Physical Domain).
Uses the reusable tools/ieee toolkit.
"""

import sys
from pathlib import Path

# Add project root to Python path so 'tools' is importable
project_root = Path(__file__).resolve().parents[3]
sys.path.insert(0, str(project_root))

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from tools.ieee.formatter import setup_ieee_document
from tools.ieee import styles


def build_physical_domain_section(output_path="Physical_Domain_Features_IEEE.docx"):
    doc = setup_ieee_document(two_column=True)

    # Section heading
    styles.add_section_heading(doc, "III", "METHODOLOGY")
    # Subsection heading
    styles.add_subsection_heading(doc, "E", "Physical-Domain Feature Engineering")

    # Intro paragraph
    intro = (
        "Để làm phong phú tập dữ liệu với các đại lượng điện năng mang tính vật lý, "
        "hai biến phái sinh quan trọng — Công suất biểu kiến (S) và Góc lệch pha (φ) — "
        "được tính toán từ các biến đo công suất thô. Những đặc trưng này giúp phơi bày "
        "các chế độ vận hành bất thường mà công suất tác dụng (P) đơn thuần không thể phát hiện, "
        "chẳng hạn như hiện tượng kẹt rô-to hoặc rò rỉ cách điện tiến triển, từ đó cung cấp "
        "các chỉ báo hữu ích cho mô hình dự báo tải và phân loại rủi ro."
    )
    styles.add_body_paragraph(doc, intro)

    # 1) Apparent Power
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("1) Công suất biểu kiến (S):")
    styles.set_run_font(run, bold=True)
    styles.set_paragraph_spacing(p, before=3, after=3)

    s_desc = (
        "Công suất biểu kiến là độ lớn của công suất phức trong mạch điện xoay chiều. "
        "Đại lượng này quyết định giới hạn chịu tải nhiệt và cơ học của dây cáp, máy biến áp "
        "và thiết bị đóng cắt. Khi công suất tác dụng không đổi nhưng công suất phản kháng tăng "
        "(do kẹt rô-to hoặc bão hòa từ trình), S sẽ phình to — đây là tín hiệu cảnh báo quá dòng."
    )
    styles.add_body_paragraph(doc, s_desc)

    styles.add_body_paragraph(doc, "Định nghĩa:", bold=True)
    styles.add_equation(doc, "S = √(P² + Q²)", "8")
    vars_s = (
        "trong đó P = Usage_kWh là công suất tác dụng (kW), và Q = Lagging_Current_Reactive.Power_kVarh "
        "là công suất phản kháng trễ (kVar)."
    )
    styles.add_body_paragraph(doc, vars_s)

    # 2) Phase Angle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("2) Góc lệch pha (φ):")
    styles.set_run_font(run, bold=True)
    styles.set_paragraph_spacing(p, before=3, after=3)

    phi_desc = (
        "Góc lệch pha định lượng độ dịch thời gian giữa sóng điện áp và dòng điện. "
        "Đại lượng này cung cấp thang đo tuyến tính sắc nét hơn Hệ số công suất (PF), "
        "đặc biệt khi phụ tải trở nên phi tuyến hoặc mang tính cảm kháng cao."
    )
    styles.add_body_paragraph(doc, phi_desc)

    styles.add_body_paragraph(doc, "Định nghĩa:", bold=True)
    styles.add_equation(doc, "φ = arccos(PF)", "9")
    vars_phi = (
        "trong đó PF = Lagging_Current_Power_Factor được giới hạn trong khoảng [0, 1] trước khi áp dụng "
        "hàm arccos nhằm đảm bảo tính ổn định số học. Kết quả φ được biểu diễn bằng radian."
    )
    styles.add_body_paragraph(doc, vars_phi)

    # Table
    styles.add_ieee_table(
        doc,
        caption_number="I",
        caption_text="CÁC ĐẶC TRƯNG MIỀN VẬT LÝ",
        headers=["Đặc trưng", "Ký hiệu", "Công thức", "Đơn vị", "Ý nghĩa vật lý"],
        rows=[
            ["Công suất biểu kiến", "S", "√(P² + Q²)", "kVA",
             "Tổng công suất đặt lên hệ thống truyền tải; cảnh báo quá tải sớm."],
            ["Góc lệch pha", "φ", "arccos(PF)", "rad",
             "Thang đo tuyến tính cho tính phi tuyến của phụ tải; phân biệt rõ hơn PF gần 1."],
        ],
    )

    # Interpretation
    styles.add_body_paragraph(doc, "Giải thích vật lý và ứng dụng:", bold=True)
    interp = (
        "Trong bối cảnh nhà máy thép, cả S và φ đóng vai trò chỉ báo nhạy với bất thường. "
        "Sự gia tăng đồng thời của S kèm theo φ tăng nhanh thường báo hiệu quá tải cảm kháng "
        "hoặc hỏng hóc cách điện. Ngược lại, S tăng trong khi φ ổn định có thể chỉ ra quá tải thuần trở. "
        "Bằng cách đưa các kênh phái sinh từ vật lý vào không gian đặc trưng, các mô hình dự báo "
        "và phân loại phía sau sẽ có thêm các giả thiết tiên nghiệm mà các đặc trưng thuần thống kê "
        "(trung bình trượt, năng lượng sóng) không thể cung cấp."
    )
    styles.add_body_paragraph(doc, interp)

    # Integration
    styles.add_body_paragraph(doc, "Tích hợp với pipeline:", bold=True)
    integ = (
        "Các đặc trưng vật lý được tính toán sau giai đoạn trích xuất miền thời gian và miền tần số, "
        "và trước bước gán nhãn bất thường. Thứ tự này đảm bảo các quy tắc bất thường dựa trên vật lý "
        "(chạy không tải, rò rỉ, quá tải) có thể trực tiếp tham chiếu đến S và φ cùng với các biến công suất gốc."
    )
    styles.add_body_paragraph(doc, integ)

    doc.save(output_path)
    print("[OK] Saved:", output_path)


if __name__ == "__main__":
    build_physical_domain_section()