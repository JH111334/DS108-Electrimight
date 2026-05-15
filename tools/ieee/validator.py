#!/usr/bin/env python3
"""
IEEE docx validator — programmatically checks if a .docx conforms to IEEE format.

Usage:
    python tools/ieee/validator.py <path_to_docx>
"""

import sys
import re
from docx import Document
from docx.shared import Inches, Pt

from tools.ieee.config import (
    FONT_NAME,
    SIZE_BODY,
    SIZE_SECTION,
    SIZE_TITLE,
    PAGE_WIDTH,
    PAGE_HEIGHT,
    LEFT_MARGIN,
    RIGHT_MARGIN,
    TOP_MARGIN,
    BOTTOM_MARGIN,
)


def _approx(val, target, tol=0.05):
    """Check if val approx equals target within tolerance."""
    return abs(val - target) < tol


def validate_page_setup(doc) -> list:
    """Check page dimensions and margins."""
    issues = []
    sec = doc.sections[0]
    pw = sec.page_width
    ph = sec.page_height
    lm = sec.left_margin
    rm = sec.right_margin
    tm = sec.top_margin
    bm = sec.bottom_margin

    if not _approx(pw, PAGE_WIDTH):
        issues.append(f"Page width: {pw.inches:.2f}\" (expected {PAGE_WIDTH.inches}\")")
    if not _approx(ph, PAGE_HEIGHT):
        issues.append(f"Page height: {ph.inches:.2f}\" (expected {PAGE_HEIGHT.inches}\")")
    if not _approx(lm, LEFT_MARGIN):
        issues.append(f"Left margin: {lm.inches:.2f}\" (expected {LEFT_MARGIN.inches}\")")
    if not _approx(rm, RIGHT_MARGIN):
        issues.append(f"Right margin: {rm.inches:.2f}\" (expected {RIGHT_MARGIN.inches}\")")
    if not _approx(tm, TOP_MARGIN):
        issues.append(f"Top margin: {tm.inches:.2f}\" (expected {TOP_MARGIN.inches}\")")
    if not _approx(bm, BOTTOM_MARGIN):
        issues.append(f"Bottom margin: {bm.inches:.2f}\" (expected {BOTTOM_MARGIN.inches}\")")
    return issues


def validate_fonts(doc) -> list:
    """Check that all runs use Times New Roman."""
    issues = []
    bad_fonts = set()
    for para in doc.paragraphs:
        for run in para.runs:
            fname = run.font.name
            if fname and fname != FONT_NAME:
                bad_fonts.add(fname)
    if bad_fonts:
        issues.append(f"Non-IEEE fonts found: {bad_fonts}. Expected: {FONT_NAME}")
    return issues


def validate_headings(doc) -> list:
    """Check heading formatting."""
    issues = []
    for para in doc.paragraphs:
        style = para.style.name
        if style == "Heading 1":
            text = para.text.strip()
            if not text.isupper():
                issues.append(f"Section heading not ALL CAPS: '{text[:50]}...'")
            if not any(para.text.startswith(r) for r in ["I.", "II.", "III.", "IV.", "V.", "REFERENCES"]):
                if para.text != "REFERENCES":
                    pass  # Allow non-Roman if it's REFERENCES
    return issues


def validate_references(doc) -> list:
    """Check reference section exists and uses [N] format."""
    issues = []
    ref_started = False
    ref_count = 0
    for para in doc.paragraphs:
        if para.style.name == "Heading 1" and "REFERENCES" in para.text.upper():
            ref_started = True
            continue
        if ref_started and para.text.strip():
            if not re.match(r"^\[\d+\]", para.text.strip()):
                issues.append(f"Reference lacks [N] prefix: '{para.text[:60]}...'")
            ref_count += 1
    if ref_count == 0:
        issues.append("No references found or reference section missing.")
    return issues


def validate(doc_path: str) -> dict:
    """Run full validation and return report."""
    doc = Document(doc_path)
    report = {
        "file": doc_path,
        "page_setup": validate_page_setup(doc),
        "fonts": validate_fonts(doc),
        "headings": validate_headings(doc),
        "references": validate_references(doc),
    }
    report["total_issues"] = sum(len(v) for v in report.values() if isinstance(v, list))
    return report


def print_report(report: dict):
    print(f"\n{'='*60}")
    print(f"IEEE Validation Report: {report['file']}")
    print(f"{'='*60}")
    for category, issues in report.items():
        if category in ("file", "total_issues"):
            continue
        status = "PASS" if not issues else f"{len(issues)} issue(s)"
        print(f"\n{category.upper().replace('_', ' ')}: {status}")
        for issue in issues:
            print(f"   - {issue}")
    print(f"\nTOTAL ISSUES: {report['total_issues']}")
    print("=" * 60)


if __name__ == "__main__":
    import re  # noqa
    if len(sys.argv) < 2:
        print("Usage: python validator.py <path_to_docx>")
        sys.exit(1)
    report = validate(sys.argv[1])
    print_report(report)
