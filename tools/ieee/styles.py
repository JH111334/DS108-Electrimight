#!/usr/bin/env python3
"""
Low-level style helpers for IEEE docx generation.
Wraps python-docx run/paragraph manipulation with IEEE constants.
"""

from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from tools.ieee.config import (
    FONT_NAME,
    SIZE_TITLE,
    SIZE_AUTHOR,
    SIZE_SECTION,
    SIZE_SUBSECTION,
    SIZE_BODY,
    SIZE_ABSTRACT,
    SIZE_KEYWORDS,
    SIZE_REFERENCE,
    SIZE_CAPTION,
    SPACE_AFTER_TITLE,
    SPACE_AFTER_AUTHOR,
    SPACE_BEFORE_SECTION,
    SPACE_AFTER_SECTION,
    SPACE_BEFORE_SUBSECTION,
    SPACE_AFTER_SUBSECTION,
    SPACE_AFTER_ABSTRACT,
    SPACE_AFTER_KEYWORDS,
    BODY_FIRST_LINE_INDENT,
    REF_HANGING_INDENT,
    LINE_SPACING_SINGLE,
)


def set_run_font(run, name=FONT_NAME, size=None, bold=False, italic=False, color=None):
    """Apply font settings to a run (including all script types)."""
    font = run.font
    font.name = name
    if size is not None:
        font.size = size
    font.bold = bold
    font.italic = italic
    if color is not None:
        font.color.rgb = color
    # XML-level font assignment for robustness
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_spacing(paragraph, before=0, after=0, line=LINE_SPACING_SINGLE):
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def set_hanging_indent(paragraph, indent=REF_HANGING_INDENT):
    """Set hanging indent for reference entries."""
    pf = paragraph.paragraph_format
    pf.left_indent = indent
    pf.first_line_indent = Cm(-indent.cm)


def add_title(doc, text):
    """Add IEEE title: 14pt bold, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=SIZE_TITLE, bold=True)
    set_paragraph_spacing(p, after=SPACE_AFTER_TITLE)
    return p


def add_author_block(doc, text):
    """Add author block: 12pt, centered."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_run_font(run, size=SIZE_AUTHOR)
    set_paragraph_spacing(p, after=SPACE_AFTER_AUTHOR)
    return p


def add_abstract(doc, text):
    """Add Abstract: 'Abstract—' bold, body 9pt, justified."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p.add_run("Abstract—")
    set_run_font(run_label, size=SIZE_ABSTRACT, bold=True)
    run_body = p.add_run(text)
    set_run_font(run_body, size=SIZE_ABSTRACT)
    set_paragraph_spacing(p, after=SPACE_AFTER_ABSTRACT)
    return p


def add_keywords(doc, terms):
    """Add Keywords: 'Index Terms—' bold, body 9pt, justified."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run_label = p.add_run("Index Terms—")
    set_run_font(run_label, size=SIZE_KEYWORDS, bold=True)
    run_body = p.add_run(terms)
    set_run_font(run_body, size=SIZE_KEYWORDS)
    set_paragraph_spacing(p, after=SPACE_AFTER_KEYWORDS)
    return p


def add_section_heading(doc, roman_num, title):
    """Add section heading: 10pt bold, ALL CAPS, centered."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{roman_num}. {title.upper()}")
    set_run_font(run, size=SIZE_SECTION, bold=True)
    set_paragraph_spacing(p, before=SPACE_BEFORE_SECTION, after=SPACE_AFTER_SECTION)
    return p


def add_subsection_heading(doc, letter, title):
    """Add subsection: 10pt bold italic, left-aligned."""
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(f"{letter}. {title}")
    set_run_font(run, size=SIZE_SUBSECTION, bold=True, italic=True)
    set_paragraph_spacing(p, before=SPACE_BEFORE_SUBSECTION, after=SPACE_AFTER_SUBSECTION)
    return p


def add_body_paragraph(doc, text, first_indent=BODY_FIRST_LINE_INDENT):
    """Add body paragraph: 10pt, justified, single spacing, first-line indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = first_indent
    run = p.add_run(text)
    set_run_font(run, size=SIZE_BODY)
    set_paragraph_spacing(p, after=0)
    return p


def add_reference_entry(doc, number, text):
    """Add reference: 9pt, [N] format, hanging indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_hanging_indent(p)
    run = p.add_run(f"[{number}] {text}")
    set_run_font(run, size=SIZE_REFERENCE)
    set_paragraph_spacing(p, after=0)
    return p


def add_table_caption(doc, number, text):
    """Add table caption: 'TABLE I' bold ALL CAPS above table."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_num = p.add_run(f"TABLE {number}")
    set_run_font(run_num, size=SIZE_CAPTION, bold=True)
    run_text = p.add_run(f"  {text}")
    set_run_font(run_text, size=SIZE_CAPTION, bold=False)
    set_paragraph_spacing(p, before=3, after=3)
    return p


def add_figure_caption(doc, number, text):
    """Add figure caption: 'Fig. 1.' bold below figure."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_num = p.add_run(f"Fig. {number}.")
    set_run_font(run_num, size=SIZE_CAPTION, bold=True)
    run_text = p.add_run(f"  {text}")
    set_run_font(run_text, size=SIZE_CAPTION, bold=False)
    set_paragraph_spacing(p, before=3, after=3)
    return p
