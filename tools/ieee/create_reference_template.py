#!/usr/bin/env python3
"""
Create an IEEE conference reference template (.docx) for Pandoc.

This template defines all styles (Normal, Heading 1-4, Title, Author,
Abstract, Keywords, Reference, Figure/Table Caption) with correct
IEEE font sizes, spacing, and page setup.

Usage:
    python tools/ieee/create_reference_template.py
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn


def ensure_font(run, name="Times New Roman"):
    """Set font name on a run for all script types."""
    font = run.font
    font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)


def set_paragraph_spacing(paragraph, before=0, after=0, line=1.0, rule=None):
    """Set uniform paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if rule is not None:
        pf.line_spacing_rule = rule


def setup_styles(doc):
    """Configure default document styles to IEEE specs."""

    # --- Page Setup ---
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Normal Style ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Set Normal style font attributes for all scripts
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Title Style ---
    style = doc.styles.add_style("Title", 1) if "Title" not in [s.name for s in doc.styles] else doc.styles["Title"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(14)
    font.bold = True
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Heading 1 (Section) ---
    style = doc.styles["Heading 1"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    font.bold = True
    font.italic = False
    font.all_caps = True
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Heading 2 (Subsection) ---
    style = doc.styles["Heading 2"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    font.bold = True
    font.italic = True
    font.all_caps = False
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(6)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Heading 3 ---
    style = doc.styles["Heading 3"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    font.bold = True
    font.italic = False
    font.all_caps = False
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.0
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Heading 4 ---
    style = doc.styles["Heading 4"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(10)
    font.bold = True
    font.italic = False
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Abstract style (custom) ---
    if "Abstract" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Abstract", 1)
    else:
        style = doc.styles["Abstract"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.0
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Keywords style (custom) ---
    if "Keywords" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Keywords", 1)
    else:
        style = doc.styles["Keywords"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_before = Pt(0)
    pf.space_after = Pt(12)
    pf.line_spacing = 1.0
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Reference style (custom) ---
    if "Reference" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Reference", 1)
    else:
        style = doc.styles["Reference"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(9)
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing = 1.0
    pf.left_indent = Cm(0.63)
    pf.first_line_indent = Cm(-0.63)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")

    # --- Caption style (custom) ---
    if "Caption" not in [s.name for s in doc.styles]:
        style = doc.styles.add_style("Caption", 1)
    else:
        style = doc.styles["Caption"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(8)
    font.bold = True
    pf = style.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rFonts.set(qn("w:cs"), "Times New Roman")
    rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_sample_content(doc):
    """Add placeholder paragraphs so Pandoc can map styles correctly."""

    p = doc.add_paragraph()
    p.style = doc.styles["Title"]
    run = p.add_run("Paper Title")
    ensure_font(run)

    p = doc.add_paragraph()
    run = p.add_run("Author Name\nAffiliation, Email")
    ensure_font(run)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.style = doc.styles["Abstract"]
    run_label = p.add_run("Abstract—")
    ensure_font(run_label)
    run_label.bold = True
    run_body = p.add_run("Abstract text goes here.")
    ensure_font(run_body)

    p = doc.add_paragraph()
    p.style = doc.styles["Keywords"]
    run_label = p.add_run("Index Terms—")
    ensure_font(run_label)
    run_label.bold = True
    run_body = p.add_run("keyword1, keyword2.")
    ensure_font(run_body)

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run("I. INTRODUCTION")
    ensure_font(run)

    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run("Body text goes here.")
    ensure_font(run)

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    run = p.add_run("A. Subsection Title")
    ensure_font(run)

    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run("Body text continues.")
    ensure_font(run)

    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    run = p.add_run("TABLE I")
    ensure_font(run)
    run = p.add_run(" Sample Table Caption")
    ensure_font(run)
    run.bold = False

    p = doc.add_paragraph()
    p.style = doc.styles["Caption"]
    run = p.add_run("Fig. 1.")
    ensure_font(run)
    run = p.add_run(" Sample figure caption.")
    ensure_font(run)
    run.bold = False

    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    run = p.add_run("REFERENCES")
    ensure_font(run)

    p = doc.add_paragraph()
    p.style = doc.styles["Reference"]
    run = p.add_run("[1] Author, \"Title,\" Journal, vol. x, pp. y, Year.")
    ensure_font(run)


def main():
    doc = Document()
    setup_styles(doc)
    add_sample_content(doc)
    output_path = r"tools\ieee\templates\ieee_reference.docx"
    doc.save(output_path)
    print(f"IEEE reference template created: {output_path}")


if __name__ == "__main__":
    main()
